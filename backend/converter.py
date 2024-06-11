import math2docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sympy import latex, pretty


def output_quadratic_functions(data):
    document = Document()

    document = Document()
    counter = 0
    for variant in range(1, data["numVariants"] + 1):
        theme = document.add_paragraph("Промежутки возрастания и убывания квадратичной функции")
        theme.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph(f"Вариант {variant}")
        p.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        for numTasks in range(1, data["numTasks"] + 1):
            p2 = document.add_paragraph(f"{numTasks}. Найти промежутки возрастания и убывания квадратичной функции ")
            math2docx.add_math(p2, latex(data["pull"][counter - 1]["function"]))
            counter += 1
        document.add_page_break()

    counter = 0
    for variant in range(1, data["numVariants"] + 1):
        theme = document.add_paragraph("Промежутки возрастания и убывания квадратичной функции")
        theme.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph(f"Вариант {variant}")
        p.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        for numTasks in range(1, data["numTasks"] + 1):
            p2 = document.add_paragraph(f"{numTasks}. Ответ: ")
            if data["pull"][counter - 1]["intervals"][0]["is_grow"]:
                p2.add_run("Промежуток возрастания ")
            else:
                p2.add_run("Промежуток убывания ")
            math2docx.add_math(p2, data["pull"][counter - 1]["intervals"][0]["latex"])
            if data["pull"][counter - 1]["intervals"][1]["is_grow"]:
                p2.add_run("Промежуток возрастания ")
            else:
                p2.add_run("Промежуток убывания ")
            math2docx.add_math(p2, data["pull"][counter - 1]["intervals"][1]["latex"])
            counter += 1
    document.save("new_file.docx")
    return "new_file.docx"

    # for element in array:
    #     p = document.add_paragraph()
    #     p.add_run("y= ")
    #     math2docx.add_math(p, element[0])
    #
    #     p = document.add_paragraph()
    #     if element[1]["is_grow"]:
    #         p.add_run("Промежуток возрастания ")
    #     else:
    #         p.add_run("Промежуток убывания ")
    #     math2docx.add_math(p, element[1]["latex"])
    #
    #     p = document.add_paragraph()
    #     if element[2]["is_grow"]:
    #         p.add_run("Промежуток возрастания ")
    #     else:
    #         p.add_run("Промежуток убывания ")
    #     math2docx.add_math(p, element[2]["latex"])
    # document.save("new_file.docx")
    # return "new_file.docx"


def output_quadratic_equations(data):
    document = Document()
    counter = 0
    for variant in range(1, data["numVariants"] + 1):
        theme = document.add_paragraph("Решение квадратных уравнений")
        theme.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph(f"Вариант {variant}")
        p.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        for numTasks in range(1, data["numTasks"] + 1):
            p2 = document.add_paragraph(f"{numTasks}. Найти корни уравнения ")
            math2docx.add_math(p2, latex(data["pull"][counter - 1]["equation"]))
            counter += 1
        document.add_page_break()
    counter = 0
    for variant in range(1, data["numVariants"] + 1):
        theme = document.add_paragraph("Решение квадратных уравнений (Ответы)")
        theme.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph(f"Вариант {variant}")
        p.align = WD_PARAGRAPH_ALIGNMENT.CENTER
        for numTasks in range(1, data["numTasks"] + 1):
            p2 = document.add_paragraph(f"{numTasks}. Ответ ")
            math2docx.add_math(p2, f'{latex(data["pull"][counter - 1]["roots"][0]).replace("sqrt", "sqrt[2]")}')
            p2.add_run(" ")
            math2docx.add_math(p2, f'{latex(data["pull"][counter - 1]["roots"][1]).replace("sqrt", "sqrt[2]")}')
            counter += 1

    document.save("new_file.docx")
    return "new_file.docx"
